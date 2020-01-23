import time
import docx
from docx import Document

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

#创建word文档
file=docx.Document()

#创建标题
file.add_heading('测试报告',0)

p =file.add_paragraph(time.strftime("%Y-%m-%d",time.localtime(time.time())))#打印读取到当前系统时间
p.add_run('这边的字体是')
p.add_run('加粗').bold =True
p.add_run('和')
p.add_run('倾斜').italic =True

#创建段落
file.add_heading('第一阶段',level=1)
file.add_paragraph('测试人员：Emmie',style='Intense Quote')
file.add_paragraph('xx模块测试',style='List Bullet')
file.add_paragraph('第一个测试对象',style='List Number')

file.add_picture('213.jpg',width=Inches(1))

table =file.add_table(rows=1 ,cols=3)
table.style='Table Grid'
hdr_cells =table.rows[0].cells
hdr_cells[0].text ='测试编号'
hdr_cells[1].text ='测试内容'
hdr_cells[2].text ='测试结果'

for item in range(6):
    row_cells =table.add_row().cells
    row_cells[0].text =str(item)
    row_cells[1].text = str(item)
    row_cells[2].text = str(item)

file.save('测试案例.docx')


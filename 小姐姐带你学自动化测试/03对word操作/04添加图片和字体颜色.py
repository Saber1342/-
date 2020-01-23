from docx.shared import Inches
from docx.shared import RGBColor
import docx

#创建word文档
file=docx.Document()

#添加段落字体颜色
file.add_paragraph().add_run('这里是字体颜色测试_1').font.color.rgb = RGBColor(167,2,25)
font_1 =file.add_paragraph().add_run('这里是字体颜色测试')
font_1.font.color.rgb =RGBColor(0,255,0)

#添加图片
file.add_picture('213.jpg',width=Inches(2))

#保存为word文档
file.save('mytest2.docx')

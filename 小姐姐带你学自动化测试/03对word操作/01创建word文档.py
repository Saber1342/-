import docx

#创建一个word文档，名为file
file = docx.Document()
#往file添加标题
file.add_heading('Mytest',0)
#添加名为paragraph段落
paragraph =file.add_paragraph('Hello word!')

#保存为word文档
file.save('mytest.docx')

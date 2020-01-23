import docx

#创建一个word文档，名为file
file = docx.Document('mytest1.docx')
#往file添加标题
file.add_heading('Mytest',0)
file.add_heading('Mytest',1)
file.add_heading('Mytest',2)
file.add_heading('Mytest',3)

#保存为word文档
file.save('mytest1.docx')
